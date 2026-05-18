from pathlib import Path
import math
import textwrap

from docx import Document
from docx.enum.section import WD_ORIENTATION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[2]
OUT_DIR = ROOT / "public" / "id"
DOCX_PATH = OUT_DIR / "Lazada_Storyboard.docx"
BUYER_PNG = OUT_DIR / "Lazada_Storyboard_Buyer_Flow.png"
SELLER_PNG = OUT_DIR / "Lazada_Storyboard_Seller_Admin_Flow.png"

W, H = 3300, 2550
BLACK = (0, 0, 0)
WHITE = (255, 255, 255)


def font(name, size):
    candidates = [
        Path("C:/Windows/Fonts") / name,
        Path("C:/Windows/Fonts/arial.ttf"),
        Path("C:/Windows/Fonts/calibri.ttf"),
    ]
    for path in candidates:
        if path.exists():
            return ImageFont.truetype(str(path), size=size)
    return ImageFont.load_default()


TITLE_FONT = font("arialbd.ttf", 70)
HEADER_FONT = font("arialbd.ttf", 34)
ITEM_FONT = font("arial.ttf", 28)
DESC_FONT = font("times.ttf", 28)


def text_size(draw, text, fnt):
    if not text:
        return 0, 0
    box = draw.textbbox((0, 0), text, font=fnt)
    return box[2] - box[0], box[3] - box[1]


def wrap_by_px(draw, text, fnt, max_width):
    words = text.split()
    lines = []
    current = ""
    for word in words:
        candidate = f"{current} {word}".strip()
        if text_size(draw, candidate, fnt)[0] <= max_width:
            current = candidate
        else:
            if current:
                lines.append(current)
            current = word
    if current:
        lines.append(current)
    return lines


def centered_text(draw, xywh, lines, fnt, line_gap=8):
    x, y, w, h = xywh
    heights = [text_size(draw, line, fnt)[1] for line in lines]
    total_h = sum(heights) + line_gap * max(0, len(lines) - 1)
    cy = y + max(0, (h - total_h) / 2)
    for line, line_h in zip(lines, heights):
        line_w, _ = text_size(draw, line, fnt)
        draw.text((x + (w - line_w) / 2, cy), line, fill=BLACK, font=fnt)
        cy += line_h + line_gap


def draw_box(draw, x, y, w, h, title, items, desc):
    border = 4
    header_h = 108
    pad = 22
    desc_lines = wrap_by_px(draw, desc, DESC_FONT, w - pad * 2)
    desc_line_h = 34
    desc_h = desc_line_h * len(desc_lines)
    info_h = max(132, h - header_h - desc_h - 74)

    draw.rectangle([x, y, x + w, y + h], outline=BLACK, width=border)
    centered_text(draw, (x, y + 8, w, header_h - 16), [title], HEADER_FONT, 0)
    draw.line([x + pad, y + header_h, x + w - pad, y + header_h], fill=BLACK, width=3)

    item_lines = []
    for item in items:
        item_lines.extend(wrap_by_px(draw, item, ITEM_FONT, w - pad * 2))
    centered_text(draw, (x + pad, y + header_h + 8, w - pad * 2, info_h - 16), item_lines, ITEM_FONT, 8)

    divider_y = y + header_h + info_h
    draw.line([x + pad, divider_y, x + w - pad, divider_y], fill=(115, 115, 115), width=2)

    desc_y = divider_y + 26
    for line in desc_lines:
        draw.text((x + pad, desc_y), line, fill=BLACK, font=DESC_FONT)
        desc_y += desc_line_h


def arrow(draw, start, end):
    sx, sy = start
    ex, ey = end
    draw.line([sx, sy, ex, ey], fill=BLACK, width=4)
    angle = math.atan2(ey - sy, ex - sx)
    size = 26
    left = (
        ex - size * math.cos(angle - math.pi / 6),
        ey - size * math.sin(angle - math.pi / 6),
    )
    right = (
        ex - size * math.cos(angle + math.pi / 6),
        ey - size * math.sin(angle + math.pi / 6),
    )
    draw.polygon([end, left, right], fill=BLACK)


def title(draw, text):
    tw, th = text_size(draw, text, TITLE_FONT)
    draw.text(((W - tw) / 2, 110), text, fill=BLACK, font=TITLE_FONT)


def make_canvas():
    img = Image.new("RGB", (W, H), WHITE)
    return img, ImageDraw.Draw(img)


def buyer_page():
    img, draw = make_canvas()
    title(draw, "Lazada Storyboard - Buyer Flow")

    boxes = {
        "home": (1100, 330, 1100, 420),
        "register": (420, 900, 980, 430),
        "login": (1900, 900, 980, 430),
        "category": (420, 1465, 980, 430),
        "product": (1900, 1465, 980, 430),
        "cart": (190, 2065, 850, 430),
        "checkout": (1225, 2065, 850, 430),
        "orders": (2260, 2065, 850, 430),
    }

    draw_box(draw, *boxes["home"], "1. HOME PAGE", [
        "Logo | Search Bar | Login | Signup | Cart",
        "Category Dropdown",
        "Hero Promotions / Flash Sale",
        "Categories / Just For You",
    ], "Landing page where users browse products, search items, view promos, and enter shopping features.")
    draw_box(draw, *boxes["register"], "2. BUYER REGISTER PAGE", [
        "First / Middle / Last Name",
        "Email | PH Phone",
        "Password / Re-password",
        "Delivery Location Map",
        "[ Create Account ]",
    ], "Buyers create an account, pin a delivery address, and gain access to buyer features.")
    draw_box(draw, *boxes["login"], "3. LOGIN PAGE", [
        "Phone or Email",
        "Password or Phone OTP",
        "Google / Facebook Login",
        "Role Redirect",
        "[ Login Button ]",
    ], "Users log in as buyer, seller, or admin and are routed to the correct area.")
    draw_box(draw, *boxes["category"], "4. CATEGORY / SEARCH PAGE", [
        "Category Sidebar",
        "Price Range Filter",
        "Rating Filter",
        "Sort: Popular | Sales | Price",
        "Product Grid",
    ], "Users browse category or search results and filter products before opening a listing.")
    draw_box(draw, *boxes["product"], "5. PRODUCT DETAIL PAGE", [
        "Image Gallery | Price | Ratings",
        "Variants | Quantity",
        "Delivery / Warranty Info",
        "[ Buy Now ] [ Add to Cart ]",
    ], "Users review product details, choose options, and add items to cart or buy immediately.")
    draw_box(draw, *boxes["cart"], "6. CART PAGE", [
        "Selected Items",
        "Quantity Controls",
        "Remove Item",
        "Order Summary",
    ], "Users manage selected items and continue to checkout.")
    draw_box(draw, *boxes["checkout"], "7. CHECKOUT PAGE", [
        "Delivery Address",
        "Items",
        "COD / GCash / Card",
        "[ Place Order ]",
    ], "Users confirm delivery, payment, and total before placing an order.")
    draw_box(draw, *boxes["orders"], "8. ORDERS / ACCOUNT PAGE", [
        "Order Status",
        "Profile",
        "Address Book",
        "Wishlist / Reviews",
    ], "Buyers track purchases and manage personal account details.")

    arrow(draw, (1650, 750), (910, 898))
    arrow(draw, (1650, 750), (2390, 898))
    arrow(draw, (1400, 1115), (1898, 1115))
    arrow(draw, (2390, 1330), (910, 1463))
    arrow(draw, (1400, 1680), (1898, 1680))
    arrow(draw, (2390, 1895), (615, 2063))
    arrow(draw, (1040, 2280), (1223, 2280))
    arrow(draw, (2075, 2280), (2258, 2280))

    img.save(BUYER_PNG, quality=95)


def seller_page():
    img, draw = make_canvas()
    title(draw, "Lazada Storyboard - Seller & Admin Flow")

    boxes = {
        "landing": (1100, 330, 1100, 420),
        "register": (420, 900, 980, 430),
        "login": (1900, 900, 980, 430),
        "setup": (420, 1465, 980, 430),
        "admin": (1900, 1465, 980, 430),
        "dashboard": (420, 2065, 980, 430),
        "live": (1900, 2065, 980, 430),
    }

    draw_box(draw, *boxes["landing"], "1. SELL ON LAZADA PAGE", [
        "Seller Center Header",
        "Benefits / Seller Support",
        "PH Phone + Password",
        "SMS Verification",
    ], "Sellers begin from the Seller Center landing page and start store registration.")
    draw_box(draw, *boxes["register"], "2. SELLER REGISTER PAGE", [
        "Phone Number",
        "New Password",
        "Voice Call / SMS Tabs",
        "[ Verify with SMS ]",
    ], "New sellers create a seller account and continue to business setup.")
    draw_box(draw, *boxes["login"], "3. SELLER LOGIN PAGE", [
        "Mobile Number / Email",
        "Password",
        "OTP / QR Options",
        "[ Login Button ]",
    ], "Existing sellers or admins sign in through Seller Center.")
    draw_box(draw, *boxes["setup"], "4. SELLER SETUP PAGE", [
        "Store Name",
        "Business Name",
        "Valid ID Type + Photo",
        "Full Name on ID / ID No.",
        "[ Submit ]",
    ], "Sellers complete identity details and submit their store for review.")
    draw_box(draw, *boxes["admin"], "5. ADMIN PANEL", [
        "Overview Metrics",
        "Pending Seller Verification",
        "Buyers / Sellers Management",
        "Logs / Listings",
    ], "Admin approves seller accounts, reviews activity, and manages users or listings.")
    draw_box(draw, *boxes["dashboard"], "6. SELLER DASHBOARD", [
        "Manage Products",
        "Orders / Promotions",
        "New Product",
        "Seller Tools",
    ], "Verified sellers manage listings and navigate Seller Center tools.")
    draw_box(draw, *boxes["live"], "7. ADD PRODUCT / LIVE LISTING", [
        "Product Images",
        "Name / Category / Brand",
        "Price / Stock",
        "Description / Shipping",
    ], "Submitted products become available in storefront catalog pages for buyers.")

    arrow(draw, (1650, 750), (910, 898))
    arrow(draw, (1650, 750), (2390, 898))
    arrow(draw, (910, 1330), (910, 1463))
    arrow(draw, (2390, 1330), (2390, 1463))
    arrow(draw, (1400, 1680), (1898, 1680))
    arrow(draw, (2390, 1895), (910, 2063))
    arrow(draw, (1400, 2280), (1898, 2280))

    img.save(SELLER_PNG, quality=95)


def make_docx():
    document = Document()
    section = document.sections[0]
    section.orientation = WD_ORIENTATION.LANDSCAPE
    section.page_width = Inches(11)
    section.page_height = Inches(8.5)
    section.top_margin = Inches(0.25)
    section.bottom_margin = Inches(0.25)
    section.left_margin = Inches(0.25)
    section.right_margin = Inches(0.25)
    section.header_distance = Inches(0.49)
    section.footer_distance = Inches(0.49)

    styles = document.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(11)
    styles["Normal"].paragraph_format.space_before = Pt(0)
    styles["Normal"].paragraph_format.space_after = Pt(0)
    styles["Normal"].paragraph_format.line_spacing = 1

    for image_path in [BUYER_PNG, SELLER_PNG]:
        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 1
        run = paragraph.add_run()
        run.add_picture(str(image_path), width=Inches(10.2))
        if image_path != SELLER_PNG:
            document.add_page_break()

    document.core_properties.title = "Lazada Storyboard"
    document.core_properties.subject = "Storyboard based on the Lazada PH React project"
    document.core_properties.author = "Codex"
    document.save(DOCX_PATH)


if __name__ == "__main__":
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    buyer_page()
    seller_page()
    make_docx()
    print(DOCX_PATH)
